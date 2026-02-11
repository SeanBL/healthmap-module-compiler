from __future__ import annotations

from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.oxml.ns import qn

from .blocks import ParagraphBlock, BulletsBlock, Block
from .raw_models import (
    RawSlide,
    RawEngage1Item,
)


# --------------------------------------------------
# Utilities (copied from stable Stage 1 logic)
# --------------------------------------------------

def normalize(text: str) -> str:
    if not text:
        return ""
    return " ".join(text.replace("\u00A0", " ").split()).strip()


def is_list_paragraph(p) -> bool:
    pPr = p._p.pPr
    if pPr is None:
        return False
    return pPr.find(qn("w:numPr")) is not None


def canonical_col_label(raw: str) -> str:
    t = normalize(raw).lower()
    if "/" in t:
        t = t.split("/", 1)[0].strip()

    if t.startswith("notes and instructions"):
        return "notes"
    if t.startswith("english text"):
        return "english"
    if t.startswith("image"):
        return "image"
    if t.startswith("button labels"):
        return "button_labels"

    return t


def is_slide_header(text: str) -> bool:
    t = text.lower()
    return (
        t.startswith("header:")
        or t.startswith("slide header")
        or t.startswith("slide (header)")
    )


# --------------------------------------------------
# Structural Extraction
# --------------------------------------------------

def extract_raw_slides(docx_path: Path) -> List[RawSlide]:
    doc = Document(str(docx_path))

    slides: List[RawSlide] = []
    slide_index = 0

    current_slide: Optional[RawSlide] = None
    col_labels: List[str] = []

    engage1_mode = False
    engage1_items: List[RawEngage1Item] = []
    pending_button_labels: List[str] = []

    for tbl in doc.tables:
        if len(tbl.rows) < 2:
            continue

        row_idx = 0

        while row_idx < len(tbl.rows):
            row = tbl.rows[row_idx]
            first_cell = normalize(row.cells[0].text)

            # -----------------------------------
            # Start new slide
            # -----------------------------------
            if is_slide_header(first_cell):

                # finalize previous slide
                if current_slide is not None:
                    if current_slide.slide_type == "engage1":
                        current_slide.engage1_items = engage1_items
                    slides.append(current_slide)

                slide_index += 1
                engage1_mode = False
                engage1_items = []
                pending_button_labels = []

                header_text = first_cell
                current_slide = RawSlide(
                    slide_id=f"slide_{slide_index:03d}",
                    header=header_text,
                    slide_type="panel",  # default
                )

                # column labels are on next row
                label_row = tbl.rows[row_idx + 1]
                col_labels = [canonical_col_label(c.text) for c in label_row.cells]

                row_idx += 2
                continue

            if current_slide is None:
                row_idx += 1
                continue

            # -----------------------------------
            # Extract row content
            # -----------------------------------
            row_data = {}

            for idx, cell in enumerate(row.cells):
                if idx >= len(col_labels):
                    continue

                label = col_labels[idx]
                if not label:
                    continue

                texts = [normalize(p.text) for p in cell.paragraphs if normalize(p.text)]
                if texts:
                    row_data[label] = texts

            notes = row_data.get("notes", [])
            english = row_data.get("english", [])
            image = row_data.get("image", [])

            # -----------------------------------
            # Slide type detection
            # -----------------------------------
            if notes:
                blob = " ".join(n.lower() for n in notes)

                if "slide type = engage 1" in blob:
                    current_slide.slide_type = "engage1"
                    engage1_mode = True

                elif "slide type = engage 2" in blob:
                    current_slide.slide_type = "engage2"

                elif "slide type = quiz" in blob:
                    current_slide.slide_type = "quiz"

            # -----------------------------------
            # Engage 1 Parsing (deterministic)
            # -----------------------------------
            if engage1_mode:

                if "button_labels" in row_data:
                    pending_button_labels.extend(row_data["button_labels"])
                    row_idx += 1
                    continue

                if english:
                    blocks: List[Block] = []
                    eng_col_idx = col_labels.index("english")
                    eng_cell = row.cells[eng_col_idx]

                    paras = [p for p in eng_cell.paragraphs if normalize(p.text)]
                    bullet_buffer: List[str] = []

                    for p in paras:
                        txt = normalize(p.text)

                        if is_list_paragraph(p):
                            bullet_buffer.append(txt)
                        else:
                            if bullet_buffer:
                                blocks.append(BulletsBlock(type="bullets", items=bullet_buffer))
                                bullet_buffer = []

                            blocks.append(ParagraphBlock(type="paragraph", text=txt))

                    if bullet_buffer:
                        blocks.append(BulletsBlock(type="bullets", items=bullet_buffer))

                    label = pending_button_labels[len(engage1_items)] if len(pending_button_labels) > len(engage1_items) else f"Button {len(engage1_items)+1}"

                    engage1_items.append(
                        RawEngage1Item(
                            label=label,
                            body=blocks,
                            image=image[0] if image else None,
                        )
                    )

                row_idx += 1
                continue

            # -----------------------------------
            # Panel Parsing
            # -----------------------------------
            if current_slide.slide_type == "panel" and english:

                blocks = current_slide.body or []

                eng_col_idx = col_labels.index("english")
                eng_cell = row.cells[eng_col_idx]

                bullet_buffer: List[str] = []

                for p in eng_cell.paragraphs:
                    txt = normalize(p.text)
                    if not txt:
                        continue

                    if is_list_paragraph(p):
                        bullet_buffer.append(txt)
                    else:
                        if bullet_buffer:
                            blocks.append(BulletsBlock(type="bullets", items=bullet_buffer))
                            bullet_buffer = []

                        blocks.append(ParagraphBlock(type="paragraph", text=txt))

                if bullet_buffer:
                    blocks.append(BulletsBlock(type="bullets", items=bullet_buffer))

                current_slide.body = blocks

            if image and current_slide.slide_type == "panel":
                current_slide.image = image[0]

            row_idx += 1

    if current_slide is not None:
        if current_slide.slide_type == "engage1":
            current_slide.engage1_items = engage1_items
        slides.append(current_slide)

    return slides
